"""
DOCX export agent. Converts a markdown report to a styled .docx file that follows
the formatting rules in data/writing_agent_instructions.md:
  - Arial 12 pt, US Letter (8.5 × 11 in), 1-inch margins
  - H1 16 pt bold, H2 14 pt bold, H3 12 pt bold
  - Chicago-style footnotes collected into a References section
  - Blockquotes rendered as indented italic text (for opening/closing quotes)

Files are saved to data/reports/<filename>.docx.
"""

import re
from pathlib import Path
from typing import Any

from agents import Agent, function_tool

_REPORTS_DIR = Path(__file__).parent.parent / "data" / "reports"


def _add_runs(paragraph, text: str, base_size_pt: int = 12) -> None:
    """Parse inline markdown (bold, italic, bold-italic, code, footnote refs) and add runs."""
    from docx.shared import Pt

    # Match: ***bold-italic***, **bold**, *italic*, `code`, [^N] footnote ref
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[\^(\d+)\]")
    last = 0

    for m in pattern.finditer(text):
        if m.start() > last:
            run = paragraph.add_run(text[last : m.start()])
            run.font.name = "Arial"
            run.font.size = Pt(base_size_pt)

        if m.group(1):
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
        else:
            # Footnote reference → superscript number
            run = paragraph.add_run(m.group(5))
            run.font.superscript = True

        if not m.group(4):
            run.font.name = "Arial"
        run.font.size = Pt(base_size_pt)
        last = m.end()

    if last < len(text):
        run = paragraph.add_run(text[last:])
        run.font.name = "Arial"
        run.font.size = Pt(base_size_pt)


def _setup_doc():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()

    # Page size and margins (US Letter, 1 inch all around)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1))

    # Normal style — cast through Any because python-docx stubs type styles[] as
    # BaseStyle, but paragraph styles expose .font and .paragraph_format at runtime
    normal: Any = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)

    # Heading styles
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        st: Any = doc.styles[f"Heading {level}"]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(4)

    return doc


def _convert_markdown_to_docx(markdown_text: str, output_path: str) -> None:
    from docx.shared import Pt, Inches

    doc = _setup_doc()

    # Pull out footnote definitions: [^N]: text
    footnote_def_pat = re.compile(r"^\[\^(\d+)\]:\s*(.+)$", re.MULTILINE)
    footnotes: dict[str, str] = {}
    for m in footnote_def_pat.finditer(markdown_text):
        footnotes[m.group(1)] = m.group(2).strip()
    clean = footnote_def_pat.sub("", markdown_text).strip()

    # Split into blocks (2+ blank lines)
    blocks = re.split(r"\n{2,}", clean)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = block.splitlines()
        first = lines[0]

        # --- Horizontal rule ---
        if re.fullmatch(r"[-_*]{3,}", first.strip()) and len(lines) == 1:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            continue

        # --- Headings ---
        if first.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            _add_runs(p, first[2:].strip(), base_size_pt=16)
            continue
        if first.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            _add_runs(p, first[3:].strip(), base_size_pt=14)
            continue
        if first.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            _add_runs(p, first[4:].strip(), base_size_pt=12)
            continue

        # --- Blockquote (opening/closing quote) ---
        if all(ln.startswith("> ") for ln in lines if ln.strip()):
            quote_text = " ".join(ln[2:] for ln in lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.name = "Arial"
            run.font.size = Pt(12)
            continue

        # --- Bullet list ---
        if all(re.match(r"^[-*]\s", ln) for ln in lines if ln.strip()):
            for ln in lines:
                if not ln.strip():
                    continue
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, re.sub(r"^[-*]\s", "", ln).strip())
            continue

        # --- Numbered list ---
        if all(re.match(r"^\d+\.\s", ln) for ln in lines if ln.strip()):
            for ln in lines:
                if not ln.strip():
                    continue
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, re.sub(r"^\d+\.\s+", "", ln).strip())
            continue

        # --- Regular paragraph (join lines) ---
        full = " ".join(lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        _add_runs(p, full)

    # --- References section from footnotes ---
    if footnotes:
        doc.add_paragraph(style="Heading 2").add_run("References").font.name = "Arial"
        for num, text in sorted(footnotes.items(), key=lambda x: int(x[0])):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)

    doc.save(output_path)


@function_tool
def export_to_docx(markdown_text: str, filename: str) -> dict[str, str]:
    """Convert a markdown report to a formatted .docx file and save it to the reports directory.

    Args:
        markdown_text: Full markdown content of the report.
        filename: Base filename without extension (e.g. 'report_20240115_143022_AI_trends').
    """
    _REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-_]", "_", filename)[:80]
    output_path = _REPORTS_DIR / f"{safe_name}.docx"

    try:
        _convert_markdown_to_docx(markdown_text, str(output_path))
        return {"status": "success", "file_path": str(output_path)}
    except Exception as e:
        return {"status": "error", "error": str(e), "file_path": ""}


INSTRUCTIONS = """You are a document export agent. You will receive a markdown research report, prefixed with a suggested filename on the first line (e.g. 'Filename: report_20240115_ai_trends').

Extract the filename from the first line, then call the export_to_docx tool with:
- markdown_text: everything after the filename line
- filename: the extracted filename

After the tool runs, report the file_path that was returned."""

doc_export_agent = Agent(
    name="DocExportAgent",
    instructions=INSTRUCTIONS,
    tools=[export_to_docx],
    model="gpt-4o-mini",
)
